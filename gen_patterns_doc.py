#!/usr/bin/env python3
"""Generate 47_Attack_Patterns.docx from verified source data."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/home/sakamuri/Documents/Cfiles/honeypot/47_Attack_Patterns.docx"

# ── colour palette (subtle, plain) ────────────────────────────────────────────
GROUP_COLORS = {
    "A": RGBColor(0xC6, 0xEF, 0xCE),   # light green
    "B": RGBColor(0xFF, 0xEB, 0x9C),   # light yellow
    "C": RGBColor(0xFF, 0xC7, 0xCE),   # light red
    "D": RGBColor(0xDD, 0xEB, 0xF7),   # light blue
    "E": RGBColor(0xE2, 0xEF, 0xDA),   # light olive
}
GROUP_LABELS = {
    "A": "Group A — Credential Brute Force",
    "B": "Group B — Reconnaissance & Scanning",
    "C": "Group C — Exploitation Attempts",
    "D": "Group D — Post-Exploitation Behaviour",
    "E": "Group E — Timing & Protocol Patterns",
}

PATTERNS = [
    # id, group, name, protocols, indicator, botnet, mitre, cve (or "")
    (1,  "A", "MIRAI_DEFAULT_CREDS",      "Telnet, SSH",          "root/xc3511 or root/vizxv in first 3 attempts",                              "Mirai",         "T1110.001", ""),
    (2,  "A", "MIRAI_ADMIN_SWEEP",         "Telnet, SSH",          "admin/admin → admin/1234 → admin/password sequentially, ≥2 attempts",        "Mirai",         "T1110.001", ""),
    (3,  "A", "MOZI_ROUTER_CREDS",         "Telnet, HTTP",         "root/root or Admin/Admin on ports 23 and 80 simultaneously",                  "Mozi",          "T1110.001", ""),
    (4,  "A", "HAJIME_SLOW_BRUTE",         "SSH",                  "<2 attempts/min, randomised credential order, long session dwell",            "Hajime",        "T1110.001", ""),
    (5,  "A", "CONTROL4_TARGETED",         "Telnet",               "Exclusively root/t0talc0ntr0l4! — single attempt per connection",             "Control4",      "T1110.001", ""),
    (6,  "A", "GAFGYT_DEFAULT",            "Telnet",               "root/888888 or root/default, inter-attempt interval <500 ms",                 "Gafgyt",        "T1110.001", ""),
    (7,  "A", "SATORI_HUAWEI",             "Telnet",               "root/Zte521 or supervisor username; often probes port 37215 first",           "Satori",        "T1110.001", ""),
    (8,  "A", "FBOT_FBXROUTER",            "HTTP, Telnet",         "supervisor/zyad1234 or telecomadmin/admintelecom combos",                     "FBot",          "T1110.001", ""),
    (9,  "A", "GENERIC_DICT_FAST",         "Telnet, SSH",          ">10 attempts/min, sequential dictionary ordering, no per-IP state",           "Generic",       "T1110.001", ""),
    (10, "A", "CREDENTIAL_STUFFING",       "HTTP",                 "POST /login with base64-encoded JSON bodies, rotating User-Agent",             "Generic",       "T1110.001", ""),
    (11, "A", "SINGLE_SHOT_DEFAULT",       "Telnet, SSH, HTTP",    "Exactly 1 attempt per connection — always admin/admin or root/root",          "Generic",       "T1110.001", ""),
    (12, "A", "SUBNET_COORDINATED",        "Telnet, SSH",          "Multiple source IPs in same /24 subnet using identical credentials",           "Coordinated",   "T1110.001", ""),

    (13, "B", "PORT_SCAN_SEQUENTIAL",      "TCP",                  "Connections to ports 22, 23, 80, 1883 within 10 s from same IP",              "Masscan/Zmap",  "T1046",     ""),
    (14, "B", "BANNER_GRAB_ONLY",          "SSH, Telnet",          "Connects, reads banner, disconnects without sending any data",                 "Shodan",        "T1046",     ""),
    (15, "B", "HTTP_FINGERPRINT",          "HTTP",                 "GET /, HEAD /, GET /favicon.ico in sequence; reads Server header",             "Generic",       "T1046",     ""),
    (16, "B", "CGI_PROBE",                 "HTTP",                 "GET requests to /cgi-bin/, /shell, /command, /cmd in single session",          "Generic",       "T1203",     ""),
    (17, "B", "MQTT_TOPIC_ENUM",           "MQTT",                 "Subscribes to # or $SYS/# immediately after CONNECT",                         "Generic",       "T1046",     ""),
    (18, "B", "UPNP_PROBE",               "HTTP",                 "GET /description.xml or GET /rootDesc.xml",                                    "Generic",       "T1046",     ""),
    (19, "B", "MITRE_T1046_NETSCAN",      "TCP",                  "SYN packets across >5 ports, no ACK completion, <100 ms between each",         "Generic",       "T1046",     ""),
    (20, "B", "SLOW_RECON",               "TCP",                  "One connection every 5–60 min from same IP over a 24-hour window",             "APT",           "T1046",     ""),

    (21, "C", "SHELLSHOCK",               "HTTP",                 '() { :;}; string in User-Agent or any header/body field',                      "Generic",       "T1059",     "CVE-2014-6271"),
    (22, "C", "DASAN_RCE",                "HTTP",                 "GET /GponForm/diag_Form?images/ with command injection in dest param",          "Generic",       "T1203",     "CVE-2018-10562"),
    (23, "C", "HUAWEI_HG532_RCE",         "HTTP",                 "POST /ctrlt/DeviceUpgrade_1 with NewStatusURL shell command injection",         "Generic/Satori","T1203",     "CVE-2017-17215"),
    (24, "C", "REALTEK_SDK_RCE",          "HTTP",                 "POST /soap.cgi UPnP SUBSCRIBE body containing shell commands",                 "Generic",       "T1203",     "CVE-2021-35395"),
    (25, "C", "BUFFER_OVERFLOW_TELNET",   "Telnet",               "Single input line >256 bytes; NOP sled pattern often present",                 "Generic",       "T1203",     ""),
    (26, "C", "DIR_TRAVERSAL_HTTP",       "HTTP",                 "../ or %2e%2e in URL path, attempts to read /etc/passwd or /proc/",            "Generic",       "T1083",     ""),
    (27, "C", "COMMAND_INJECTION_HTTP",   "HTTP",                 ";, |, && or backtick in form field values or GET parameters",                  "Generic",       "T1059.004", ""),
    (28, "C", "MQTT_MALICIOUS_PUBLISH",   "MQTT",                 "PUBLISH to cmnd/ or zigbee2mqtt/ topics with shell payload in body",           "Generic",       "T1499",     ""),
    (29, "C", "LOG4SHELL_PROBE",          "HTTP",                 "${jndi:ldap:// string in any HTTP header or body field",                       "Generic",       "T1203",     "CVE-2021-44228"),
    (30, "C", "SPRING4SHELL_PROBE",       "HTTP",                 "POST with class.module.classLoader parameter",                                 "Generic",       "T1203",     "CVE-2022-22965"),

    (31, "D", "WGET_DROPPER",             "Telnet, SSH",          "wget http:// or curl http:// executed post-auth",                              "Mirai",         "T1105",     ""),
    (32, "D", "BUSYBOX_WGET_CHAIN",       "Telnet",               "/bin/busybox wget + SATORI or ECCHI keyword suffix",                           "Satori/Ecchi",  "T1105",     ""),
    (33, "D", "CHMOD_EXECUTE",            "Telnet, SSH",          "chmod +x on a file followed immediately by ./ execution",                     "Generic",       "T1059",     ""),
    (34, "D", "CRONTAB_PERSISTENCE",      "Telnet, SSH",          "crontab -e or echo >> /etc/cron* to establish persistence",                   "APT",           "T1053",     ""),
    (35, "D", "IPTABLES_MANIPULATION",    "Telnet, SSH",          "iptables -F (flush) or iptables -A (add rule) commands",                      "Generic",       "T1562",     ""),
    (36, "D", "CRYPTO_MINER_INSTALL",     "Telnet, SSH",          "xmrig, minerd, or mining pool URL (pool.) in command",                        "Cryptojacker",  "T1496",     ""),
    (37, "D", "C2_CALLBACK_ATTEMPT",      "Telnet, SSH",          "nc, ncat, or bash -i >& /dev/tcp/ reverse shell commands",                   "APT",           "T1059",     ""),
    (38, "D", "SELF_PROPAGATION",         "Telnet, SSH",          "for i in loop scanning subnet with embedded telnet commands",                  "Worm",          "T1210",     ""),
    (39, "D", "LOG_WIPE",                 "Telnet, SSH",          "rm /var/log/, echo > /var/log/messages, or history -c",                      "APT",           "T1070",     ""),

    (40, "E", "HAWKES_BURST_A",           "Any",                  ">20 events/min burst then silence >30 min (Hawkes cluster type A)",            "Generic",       "T1498",     ""),
    (41, "E", "HAWKES_BURST_B",           "Any",                  "5–15 events/min sustained >2 hours (Hawkes cluster type B)",                   "Generic",       "T1498",     ""),
    (42, "E", "HAWKES_PERIODIC",          "Any",                  "Regular inter-arrival time — coefficient of variation < 0.30",                 "Generic",       "T1498",     ""),
    (43, "E", "DIURNAL_NIGHT",            "Any",                  "≥70% of events from a source IP fall between 00:00–06:00 UTC",                "Generic",       "T1498",     ""),
    (44, "E", "MULTI_PROTOCOL_CHAIN",     "Telnet, HTTP, MQTT",   "Same source IP hits Telnet + HTTP + MQTT within a 60 s window",               "Generic",       "T1046",     ""),
    (45, "E", "TLS_DOWNGRADE",            "SSH, HTTPS",           "Weak cipher suites offered: RC4-SHA, 3DES, NULL, EXPORT-RC4-40",              "Generic",       "T1562",     ""),
    (46, "E", "MQTT_QOS_ABUSE",           "MQTT",                 "Rapid QoS 2 PUBLISH messages to exhaust broker state table",                  "Generic",       "T1499",     ""),
    (47, "E", "ZERO_DAY_ANOMALY",         "Any",                  "Anomaly score > threshold; no pattern 1–46 matched — DBSCAN outlier",         "Unknown",       "T1203",     ""),
]

CLASSIFIER_LAYER = {
    "A": "On-device (ESP32 firmware)",
    "B": "On-device (ESP32 firmware)",
    "C": "Backend classifier only",
    "D": "On-device (ESP32 firmware)",
    "E": "Backend classifier only",
}


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def rgb_to_hex(r: RGBColor) -> str:
    return f"{r[0]:02X}{r[1]:02X}{r[2]:02X}"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)


def build():
    doc = Document()

    # ── Title ─────────────────────────────────────────────────────────────────
    title = doc.add_heading("47 Attack Pattern Taxonomy", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("ESP32 IoT Honeypot Project")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True

    doc.add_paragraph()

    # ── Dataset source ────────────────────────────────────────────────────────
    add_heading(doc, "Dataset / Source", level=1)

    add_para(doc,
        "The 47-pattern taxonomy is derived from the following primary research paper:")

    p = doc.add_paragraph()
    r = p.add_run(
        '"Using Honeypots to Model Botnet Attacks on the Internet of Medical Things"')
    r.bold = True
    r.italic = True
    p.add_run(
        "  —  PubMed Central, PMC9264116.")

    add_para(doc,
        "The paper processed 128,737 real attack sessions captured by IoT honeypots and "
        "resolved them into 47 distinct botnet attack clusters using a Multivariate "
        "Hawkes Process model, further refined into 121 temporal sub-clusters.")

    add_para(doc, "Key findings that directly drive this implementation:")

    bullets = [
        "Weak password exploitation is the primary IoT attack surface — credential brute-force dominates (Groups A).",
        "Bots on the same /24 subnet tend to try identical credential combinations (Pattern 12: SUBNET_COORDINATED).",
        "74% of bots are segmented into fewer than 10 temporal fragments by control period.",
        "Attack timing patterns (inter-arrival rates, burst shapes) are sufficient to fingerprint botnet families (Group E).",
    ]
    for b in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(b)
        p.paragraph_format.space_after = Pt(2)

    add_para(doc,
        "Additional pattern names and indicators for Groups B, C, and D are drawn from "
        "supplementary honeypot literature covering Mirai, Mozi, Hajime, HoneyIoT, and "
        "autoencoder clustering papers, as documented in PHASE2.md.")

    doc.add_paragraph()

    # ── Classifier layer note ─────────────────────────────────────────────────
    add_heading(doc, "Where Each Group Is Classified", level=1)

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Group", "Patterns", "Classified By"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(hdr[i], "D9D9D9")

    layer_rows = [
        ("A  —  Credential Brute Force",       "1 – 12",  "On-device  (ESP32 firmware, attack_patterns.h)"),
        ("B  —  Reconnaissance & Scanning",    "13 – 20", "On-device  (ESP32 firmware, attack_patterns.h)"),
        ("C  —  Exploitation Attempts",        "21 – 30", "Backend classifier only  (heuristic_rules.py)"),
        ("D  —  Post-Exploitation Behaviour",  "31 – 39", "On-device  (ESP32 firmware, attack_patterns.h)"),
        ("E  —  Timing & Protocol Patterns",   "40 – 47", "Backend classifier only  (hawkes_classifier.py, anomaly_detector.py)"),
    ]
    for grp, pats, layer in layer_rows:
        row = tbl.add_row().cells
        row[0].text = grp
        row[1].text = pats
        row[2].text = layer

    doc.add_paragraph()

    # ── Per-group pattern tables ───────────────────────────────────────────────
    current_group = None
    for (pid, grp, name, protos, indicator, botnet, mitre, cve) in PATTERNS:
        if grp != current_group:
            current_group = grp
            add_heading(doc, GROUP_LABELS[grp], level=1)

            # Column headers
            tbl = doc.add_table(rows=1, cols=6)
            tbl.style = "Table Grid"
            hdr = tbl.rows[0].cells
            cols = ["#", "Pattern Name", "Protocol(s)", "Primary Indicator", "Botnet Family", "MITRE"]
            for i, h in enumerate(cols):
                hdr[i].text = h
                hdr[i].paragraphs[0].runs[0].bold = True
                set_cell_bg(hdr[i], "D9D9D9")

        row_cells = tbl.add_row().cells
        hex_bg = rgb_to_hex(GROUP_COLORS[grp])

        row_cells[0].text = str(pid)
        row_cells[1].text = name
        row_cells[2].text = protos
        indicator_text = indicator if not cve else f"{indicator}  [{cve}]"
        row_cells[3].text = indicator_text
        row_cells[4].text = botnet
        row_cells[5].text = mitre

        for c in row_cells:
            set_cell_bg(c, hex_bg)
            for para in c.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)

    doc.add_paragraph()

    # ── Quick reference (all 47 in one table) ─────────────────────────────────
    add_heading(doc, "Quick Reference — All 47 Patterns", level=1)

    ref_tbl = doc.add_table(rows=1, cols=5)
    ref_tbl.style = "Table Grid"
    hdr = ref_tbl.rows[0].cells
    for i, h in enumerate(["#", "Group", "Pattern Name", "Botnet Family", "MITRE"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(hdr[i], "D9D9D9")

    for (pid, grp, name, protos, indicator, botnet, mitre, cve) in PATTERNS:
        row = ref_tbl.add_row().cells
        row[0].text = str(pid)
        row[1].text = grp
        row[2].text = name
        row[3].text = botnet
        row[4].text = mitre
        hex_bg = rgb_to_hex(GROUP_COLORS[grp])
        for c in row:
            set_cell_bg(c, hex_bg)
            for para in c.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)

    doc.save(OUT)
    print(f"Saved → {OUT}")


if __name__ == "__main__":
    build()
