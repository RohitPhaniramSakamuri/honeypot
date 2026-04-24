#!/usr/bin/env python3
"""
Generate CAT-2 document for the ESP32 IoT Honeypot project.
Produces: cat2_report.docx  +  3 diagram PNGs in /tmp/
"""

import os
import textwrap
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyArrowPatch
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = "/home/sakamuri/Documents/Cfiles/honeypot"
IMG_DIR = "/tmp/cat2_imgs"
os.makedirs(IMG_DIR, exist_ok=True)


# ─────────────────────────────────────────────────────────────────────────────
# Diagram 1 — FSM (Telnet honeypot state machine)
# ─────────────────────────────────────────────────────────────────────────────
def draw_fsm():
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis("off")
    ax.set_facecolor("white")
    fig.patch.set_facecolor("white")

    states = {
        "IDLE":          (1.2, 3.0),
        "CONNECT":       (3.2, 3.0),
        "STATE_LOGIN":   (5.2, 4.5),
        "STATE_PASSWORD":(5.2, 3.0),
        "STATE_SHELL":   (5.2, 1.5),
        "DISCONNECT":    (8.2, 3.0),
    }

    node_r = 0.55
    colors = {
        "IDLE":           "#aed6f1",
        "CONNECT":        "#a9dfbf",
        "STATE_LOGIN":    "#f9e79f",
        "STATE_PASSWORD": "#f9e79f",
        "STATE_SHELL":    "#f1948a",
        "DISCONNECT":     "#d7dbdd",
    }

    for name, (cx, cy) in states.items():
        circle = plt.Circle((cx, cy), node_r, color=colors[name],
                             ec="black", lw=1.5, zorder=3)
        ax.add_patch(circle)
        ax.text(cx, cy, name.replace("_", "\n"), ha="center", va="center",
                fontsize=7.5, fontweight="bold", zorder=4)

    def arrow(src, dst, label="", rad=0.0, lx=0, ly=0):
        sx, sy = states[src]
        dx, dy = states[dst]
        ax.annotate("", xy=(dx, dy), xytext=(sx, sy),
                    arrowprops=dict(arrowstyle="->", color="black",
                                   lw=1.4,
                                   connectionstyle=f"arc3,rad={rad}"))
        mx, my = (sx + dx) / 2 + lx, (sy + dy) / 2 + ly
        if label:
            ax.text(mx, my, label, fontsize=7, color="#1a5276",
                    ha="center", va="center",
                    bbox=dict(boxstyle="round,pad=0.2", fc="white", ec="none", alpha=0.8))

    arrow("IDLE",           "CONNECT",        "TCP SYN / log connect")
    arrow("CONNECT",        "STATE_LOGIN",    "send BusyBox banner", rad=0.3, ly=0.3)
    arrow("STATE_LOGIN",    "STATE_PASSWORD", "username received",  lx=0.9)
    arrow("STATE_PASSWORD", "STATE_SHELL",    "Mirai cred match →\nauth_success", lx=0.9)
    arrow("STATE_PASSWORD", "STATE_LOGIN",    "cred mismatch →\nauth_attempt", rad=-0.35, lx=-1.1)
    arrow("STATE_SHELL",    "DISCONNECT",     "exit / timeout /\nconnection drop", rad=-0.3, ly=-0.3)
    arrow("STATE_LOGIN",    "DISCONNECT",     "timeout", rad=0.25, ly=0.4)
    arrow("STATE_SHELL",    "STATE_SHELL",    "command →\nlog + respond",
          rad=0.5)
    # self-loop label
    ax.text(6.5, 1.5, "cmd loop", fontsize=7, color="#1a5276", ha="center")

    ax.set_title("Figure 1 — Telnet Honeypot Finite State Machine (FSM)",
                 fontsize=11, fontweight="bold", pad=10)

    path = os.path.join(IMG_DIR, "fsm.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    print(f"  FSM saved → {path}")
    return path


# ─────────────────────────────────────────────────────────────────────────────
# Diagram 2 — CDFG (Control / Data Flow Graph of the full pipeline)
# ─────────────────────────────────────────────────────────────────────────────
def draw_cdfg():
    fig, ax = plt.subplots(figsize=(12, 7))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 7)
    ax.axis("off")
    ax.set_facecolor("white")
    fig.patch.set_facecolor("white")

    boxes = {
        # (label,   x,   y,   w,   h,  fill)
        "Attacker":         (0.3, 5.8, 1.6, 0.6, "#fadbd8"),
        "Telnet :23":       (0.3, 4.7, 1.6, 0.5, "#fdebd0"),
        "SSH :22":          (0.3, 4.0, 1.6, 0.5, "#fdebd0"),
        "HTTP :80":         (0.3, 3.3, 1.6, 0.5, "#fdebd0"),
        "MQTT-HP :1883":    (0.3, 2.6, 1.6, 0.5, "#fdebd0"),
        "EventLogger":      (2.7, 3.7, 1.8, 0.6, "#d5f5e3"),
        "MQTT Publish":     (2.7, 2.6, 1.8, 0.5, "#d5f5e3"),
        "Mosquitto\n:1883": (5.2, 3.2, 1.8, 0.6, "#d6eaf8"),
        "Telegraf":         (7.4, 3.2, 1.6, 0.6, "#e8daef"),
        "InfluxDB\n:8086":  (9.2, 3.2, 1.6, 0.6, "#fef9e7"),
        "Classifier\n:8000":(9.2, 1.5, 1.6, 0.6, "#fdf2f8"),
        "Grafana\n:3000":   (9.2, 4.9, 1.6, 0.6, "#eafaf1"),
    }

    def box_center(name):
        x, y, w, h, _ = boxes[name]
        return x + w / 2, y + h / 2

    for name, (x, y, w, h, color) in boxes.items():
        rect = mpatches.FancyBboxPatch((x, y), w, h,
                                       boxstyle="round,pad=0.08",
                                       facecolor=color, edgecolor="black", lw=1.2)
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2, name,
                ha="center", va="center", fontsize=8, fontweight="bold")

    def arr(src, dst, label="", color="black", rad=0.0):
        sx, sy = box_center(src)
        dx, dy = box_center(dst)
        ax.annotate("", xy=(dx, dy), xytext=(sx, sy),
                    arrowprops=dict(arrowstyle="->", color=color, lw=1.3,
                                   connectionstyle=f"arc3,rad={rad}"))
        if label:
            mx, my = (sx + dx) / 2, (sy + dy) / 2
            ax.text(mx, my, label, fontsize=6.5, color="#2c3e50",
                    ha="center", va="bottom",
                    bbox=dict(boxstyle="round,pad=0.15", fc="white", ec="none", alpha=0.85))

    # Attacker → services
    for svc in ["Telnet :23", "SSH :22", "HTTP :80", "MQTT-HP :1883"]:
        arr("Attacker", svc, color="#c0392b")

    # Services → EventLogger
    for svc in ["Telnet :23", "SSH :22", "HTTP :80", "MQTT-HP :1883"]:
        arr(svc, "EventLogger", color="#27ae60")

    arr("EventLogger",    "MQTT Publish",   "enqueue / flush",  color="#27ae60")
    arr("MQTT Publish",   "Mosquitto\n:1883","JSON / Protobuf", color="#2980b9")
    arr("Mosquitto\n:1883","Telegraf",       "subscribe #",      color="#8e44ad")
    arr("Telegraf",       "InfluxDB\n:8086", "write points",     color="#d4ac0d")
    arr("InfluxDB\n:8086","Classifier\n:8000","poll 60s",       color="#cb4335", rad=0.15)
    arr("Classifier\n:8000","InfluxDB\n:8086","write enriched",  color="#cb4335", rad=0.15)
    arr("InfluxDB\n:8086","Grafana\n:3000",  "Flux query",       color="#1abc9c")

    # Data labels
    ax.text(5.5, 0.4, "Control flow (arrows)  |  Data flow (arrow labels)  |  All services on honeypot_net Docker bridge",
            fontsize=7, ha="center", color="#555")

    ax.set_title("Figure 2 — System CDFG: Control & Data Flow Graph",
                 fontsize=11, fontweight="bold", pad=10)

    path = os.path.join(IMG_DIR, "cdfg.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    print(f"  CDFG saved → {path}")
    return path


# ─────────────────────────────────────────────────────────────────────────────
# Diagram 3 — DES (Discrete Event Simulation timeline)
# ─────────────────────────────────────────────────────────────────────────────
def draw_des():
    fig, ax = plt.subplots(figsize=(12, 5))
    ax.set_xlim(-0.5, 13)
    ax.set_ylim(-0.5, 5.5)
    ax.axis("off")
    ax.set_facecolor("white")
    fig.patch.set_facecolor("white")

    # Swimlane rows
    lanes = ["Attacker", "TCP Stack", "EventLogger\n(queue)", "MQTT / Broker", "InfluxDB"]
    lane_y = {name: 4.5 - i * 1.0 for i, name in enumerate(lanes)}
    lane_colors = ["#fadbd8", "#fef9e7", "#d5f5e3", "#d6eaf8", "#f5eef8"]

    for i, (name, y) in enumerate(lane_y.items()):
        ax.barh(y, 13, height=0.7, left=-0.5, color=lane_colors[i],
                edgecolor="#ccc", linewidth=0.6, zorder=1)
        ax.text(-0.4, y, name, va="center", ha="left", fontsize=8, fontweight="bold")

    # Timeline axis
    ax.axhline(-0.3, color="black", lw=1)
    for t in range(0, 13):
        ax.plot([t, t], [-0.4, -0.2], color="black", lw=0.8)
        if t % 2 == 0:
            ax.text(t, -0.55, f"t+{t}s", ha="center", fontsize=7)
    ax.text(13, -0.55, "time →", ha="right", fontsize=8)

    # Events
    events = [
        # (t,  lane,                label,         color)
        (0.0,  "Attacker",          "TCP SYN :23", "#e74c3c"),
        (0.1,  "TCP Stack",         "CONNECT evt", "#2980b9"),
        (0.2,  "EventLogger\n(queue)", "enqueue\nconnect", "#27ae60"),
        (1.0,  "Attacker",          "send user",   "#e74c3c"),
        (1.1,  "TCP Stack",         "auth_attempt\nevt",   "#2980b9"),
        (1.2,  "EventLogger\n(queue)", "enqueue\nauth",    "#27ae60"),
        (1.5,  "Attacker",          "Mirai cred\nmatch",   "#c0392b"),
        (1.6,  "TCP Stack",         "auth_success\n→ immediate publish", "#e67e22"),
        (1.7,  "MQTT / Broker",     "publish\nauth topic", "#8e44ad"),
        (2.5,  "Attacker",          "wget cmd",    "#e74c3c"),
        (2.6,  "TCP Stack",         "exploit evt", "#c0392b"),
        (2.7,  "EventLogger\n(queue)", "enqueue\nexploit", "#27ae60"),
        (2.8,  "MQTT / Broker",     "publish\nexploit",    "#8e44ad"),
        (2.9,  "InfluxDB",          "write point", "#1abc9c"),
        (5.0,  "Attacker",          "disconnect",  "#95a5a6"),
        (10.0, "EventLogger\n(queue)", "30s flush\n(heartbeat)", "#16a085"),
        (10.1, "MQTT / Broker",     "heartbeat\npublish",  "#8e44ad"),
        (10.2, "InfluxDB",          "batch write", "#1abc9c"),
    ]

    for t, lane, label, color in events:
        y = lane_y[lane]
        ax.plot(t, y, "o", color=color, ms=7, zorder=5)
        ax.text(t, y + 0.28, label, ha="center", va="bottom",
                fontsize=6.2, color=color, zorder=6,
                bbox=dict(boxstyle="round,pad=0.1", fc="white", ec="none", alpha=0.7))
        ax.vlines(t, -0.3, y, colors=color, lw=0.6, linestyles="dashed", alpha=0.4, zorder=2)

    ax.set_title("Figure 3 — Discrete Event Simulation (DES): single Telnet session timeline",
                 fontsize=11, fontweight="bold", pad=10)

    path = os.path.join(IMG_DIR, "des.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    print(f"  DES saved → {path}")
    return path


# ─────────────────────────────────────────────────────────────────────────────
# Document builder
# ─────────────────────────────────────────────────────────────────────────────
def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_para(doc, text, bold_first=None):
    p = doc.add_paragraph()
    if bold_first:
        run = p.add_run(bold_first)
        run.bold = True
        p.add_run(" " + text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix + ": ")
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(2)

def add_image(doc, path, caption, width=Inches(5.5)):
    doc.add_picture(path, width=width)
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(9)
    doc.add_paragraph()  # spacer


def build_doc(fsm_path, cdfg_path, des_path):
    doc = Document()

    # Title
    title = doc.add_heading("ESP32 IoT Honeypot — CAT-2 Evaluation Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ── Q1 ───────────────────────────────────────────────────────────────────
    add_heading(doc, "1. FSM, CDFG, and DES — Course Project")

    add_heading(doc, "Finite State Machine (FSM)", level=2)
    add_para(doc,
        "The Telnet honeypot service is the most stateful component of the firmware. "
        "Each connected client progresses through three explicit states managed in "
        "telnet_honeypot.cpp. The FSM below captures every transition, the event that "
        "triggers it, and the side-effect logged.")
    add_image(doc, fsm_path,
              "Figure 1 — Telnet Honeypot FSM. "
              "States: IDLE → CONNECT → STATE_LOGIN → STATE_PASSWORD → STATE_SHELL → DISCONNECT. "
              "A Mirai credential match causes AUTH_SUCCESS and shell escalation; "
              "any other credential cycles back to STATE_LOGIN.")

    add_para(doc,
        "STATE_SHELL is a self-looping state: every command line received is logged as a "
        "'command' or 'exploit' event (exploit if the command starts with wget/curl or "
        "/bin/busybox) before the shell prompt is re-issued. "
        "The SSH and HTTP services follow simpler two-state machines "
        "(CONNECT → DISCONNECT) because they do not maintain persistent sessions — "
        "the SSH honeypot reads the banner and closes, and HTTP is stateless by design.")

    add_heading(doc, "Control and Data Flow Graph (CDFG)", level=2)
    add_para(doc,
        "The CDFG describes how data (attack events) flows from origin (attacker) through "
        "each processing node, and how control decisions (pattern classification, immediate "
        "vs. batched publish) alter that flow.")
    add_image(doc, cdfg_path,
              "Figure 2 — System CDFG. "
              "Four protocol honeypots feed into a shared EventLogger. "
              "The EventLogger serialises events as JSON (consumed by Telegraf → InfluxDB) "
              "and as Protobuf on a parallel topic (honeypot/events/pb/*). "
              "The classifier microservice polls InfluxDB every 60 s, enriches records, "
              "and writes them back. Grafana queries InfluxDB via Flux.")
    add_para(doc,
        "A key control decision in the CDFG is the 'immediate publish' branch: "
        "auth_success, command, and exploit events bypass the 30-second batch buffer "
        "and are published at once (EventLogger::_isImmediate). "
        "All other events are queued in a 50-slot circular buffer and flushed on the "
        "heartbeat tick. This prevents network congestion from high-frequency scans "
        "while still guaranteeing low latency for high-value events.")

    add_heading(doc, "Discrete Event Simulation (DES)", level=2)
    add_para(doc,
        "The honeypot system is entirely event-driven. There are no blocking waits in the "
        "main loop — all I/O is handled via AsyncTCP callbacks and WiFiServer polling. "
        "The DES timeline below traces a single Telnet session from TCP connection to "
        "heartbeat flush, showing which component processes each event and at what "
        "relative time offset.")
    add_image(doc, des_path,
              "Figure 3 — DES timeline for a Mirai Telnet session. "
              "The attacker connects (t=0 s), sends credentials (t=1 s), receives shell "
              "access (t=1.5 s), executes a wget dropper (t=2.5 s), and disconnects (t=5 s). "
              "The EventLogger flushes its queue on the 30-second heartbeat (t=10 s). "
              "Immediate-publish events (auth_success, exploit) hit InfluxDB within ~0.3 s "
              "of the triggering packet.")

    doc.add_paragraph()

    # ── Q2 ───────────────────────────────────────────────────────────────────
    add_heading(doc, "2. Code Optimizations — Course Project")

    add_heading(doc, "Trivial Techniques", level=2)
    add_bullet(doc,
        "All per-event buffers (proto, src_ip, username, password, command, event_type) "
        "are fixed-size char arrays inside HoneypotEvent, declared as a static array of "
        "50 elements. No heap allocation occurs at runtime for the event queue — this "
        "eliminates malloc overhead and fragmentation on a 520 KB SRAM device.",
        "Static allocation")
    add_bullet(doc,
        "memset(&ev, 0, sizeof(ev)) zeroes the struct in one call before populating fields, "
        "avoiding partial initialisation bugs and ensuring unused fields default to empty "
        "strings.",
        "Bulk zero-init")
    add_bullet(doc,
        "strncpy with explicit size - 1 and a manual null-terminator write is used throughout "
        "instead of strcpy, preventing buffer overflows when attacker input exceeds the "
        "field size.",
        "Bounded string copy")

    add_heading(doc, "Levels of Programming", level=2)
    add_bullet(doc,
        "ESPAsyncWebServer, PubSubClient, and LiquidCrystal_I2C are used at the library API "
        "level. These abstractions handle protocol framing, ACK management, and I2C "
        "communication, keeping application code concise.",
        "API level")
    add_bullet(doc,
        "ESP.getFreeHeap() and ESP.getCpuFreqMHz() are Arduino HAL wrappers over SFRs "
        "(Special Function Registers). The GPIO calls (pinMode, digitalWrite, esp32-hal-gpio) "
        "operate one level above bare-metal SFR writes but below a full OS driver.",
        "SFR / HAL level")
    add_bullet(doc,
        "The firmware does not contain hand-written assembly. All performance-critical paths "
        "(string scanning in isMiraiCred, Protobuf encoding via nanopb) compile down to "
        "optimised ARM Thumb-2 through the GCC toolchain with -O2 (PlatformIO default).",
        "Assembly level (via compiler)")
    add_bullet(doc,
        "The circular buffer enqueue/dequeue is O(1) time and O(N) space where N=50. "
        "isMiraiCred performs a linear scan over the credential table "
        "(O(M) where M ≈ 100 pairs); acceptable given the small table size and infrequent "
        "call rate.",
        "Time and space complexity")

    add_heading(doc, "Cache and Pipeline Efficiency", level=2)
    add_bullet(doc,
        "HoneypotEvent is a flat struct with no pointers, so the entire 50-element queue "
        "fits in a contiguous memory block. Sequential access during flushQueue() walks "
        "this block linearly, which is cache-friendly on the ESP32's Harvard-architecture "
        "data cache.",
        "Struct layout")
    add_bullet(doc,
        "StaticJsonDocument<512> is stack-allocated in _publishEvent, avoiding a heap "
        "allocation per publish and reducing cache misses compared to DynamicJsonDocument.",
        "Stack-allocated JSON")
    add_bullet(doc,
        "The Telnet loop reads one byte at a time in a while(client.available()) loop "
        "rather than blocking on read(), keeping the CPU free for other clients and the "
        "MQTT keepalive between characters.",
        "Non-blocking I/O")

    add_heading(doc, "Algorithm Implementation", level=2)
    add_bullet(doc,
        "The event queue is a power-of-two-sized circular buffer using head/tail indices. "
        "On overflow the oldest entry is silently dropped (head advances), which is the "
        "correct trade-off for a resource-constrained device where losing old data is "
        "preferable to a crash.",
        "Circular buffer (O(1))")
    add_bullet(doc,
        "nanopb encodes HoneypotEvent structs to Protobuf binary, which is approximately "
        "67% smaller than the equivalent JSON payload (measured in the serial log: typical "
        "event ~180 bytes JSON vs ~60 bytes Protobuf). This reduces MQTT publish time and "
        "Wi-Fi radio-on duration.",
        "Protobuf serialisation")
    add_bullet(doc,
        "MQTT exponential backoff on publish failure (1 s → 2 s → 4 s → max 30 s) "
        "prevents the firmware from hammering a unavailable broker and draining the battery.",
        "Exponential backoff")

    add_heading(doc, "Energy Efficiency", level=2)
    add_bullet(doc,
        "All four protocol services use AsyncTCP, which is interrupt/callback-driven. "
        "The CPU is not spinning in a blocking accept() loop; it executes code only when "
        "a packet arrives. Between events the ESP32 runs its FreeRTOS idle task, which "
        "invokes the light-sleep automatic light sleep feature.",
        "Interrupt-driven I/O (AsyncTCP)")
    add_bullet(doc,
        "Events are batched and flushed every 30 seconds rather than being sent "
        "individually. This reduces the number of times the Wi-Fi radio must enter the "
        "TX-active state, which is the dominant power consumer on the ESP32 "
        "(≈170 mA TX vs ≈20 mA modem-sleep).",
        "Batched MQTT publish")
    add_bullet(doc,
        "The LED on GPIO 18 is only driven HIGH when a Telnet client is currently "
        "connected (TelnetHoneypot::isOccupied()), not on a timer. This avoids "
        "unnecessary GPIO switching energy.",
        "Event-driven GPIO")

    doc.add_paragraph()

    # ── Q3 ───────────────────────────────────────────────────────────────────
    add_heading(doc, "3. Programming Tools — Course Project")

    add_heading(doc, "Firmware Build Toolchain", level=2)
    add_bullet(doc,
        "Used as the build system and package manager for the ESP32 firmware. "
        "platformio.ini declares the target board (esp32doit-devkit-v1), framework "
        "(Arduino), monitor speed (115200), and all library dependencies. "
        "Commands: pio run (compile), pio run --target upload (flash), "
        "pio device monitor (serial console).",
        "PlatformIO CLI")
    add_bullet(doc,
        "The Espressif esp32 platform (via PlatformIO) bundles the Xtensa GCC cross-compiler "
        "(xtensa-esp32-elf-g++). Build flags -DCORE_DEBUG_LEVEL=3 enable verbose ESP-IDF "
        "log output over UART.",
        "GCC / Xtensa toolchain")
    add_bullet(doc,
        "nanopb is a C implementation of Protocol Buffers targeting embedded systems. "
        "The .proto schema (src/honeypot_event.proto) is compiled offline with "
        "protoc + nanopb_generator to produce honeypot_event.pb.c/.h. "
        "The generated code is checked in so no protoc invocation is needed at build time.",
        "nanopb (Protocol Buffers)")

    add_heading(doc, "Libraries Used in Firmware", level=2)
    add_bullet(doc, "knolleary/PubSubClient ^2.8 — MQTT client, handles broker connection, "
               "keepalive, and publish/subscribe over WiFiClient.", "PubSubClient")
    add_bullet(doc, "bblanchon/ArduinoJson ^6.21.3 — JSON serialisation for MQTT payloads. "
               "StaticJsonDocument avoids heap allocation.", "ArduinoJson")
    add_bullet(doc, "esphome/AsyncTCP-esphome ^2.1.1 — non-blocking TCP server and client "
               "built on lwIP callbacks. Used by all four protocol honeypots.", "AsyncTCP")
    add_bullet(doc, "ottowinter/ESPAsyncWebServer-esphome ^3.1.0 — asynchronous HTTP server, "
               "handles GET/POST routing and 404 catch-all.", "ESPAsyncWebServer")
    add_bullet(doc, "marcoschwartz/LiquidCrystal_I2C ^1.1.4 — I2C driver for the 16x2 LCD, "
               "used to display the assigned IP address on boot.", "LiquidCrystal_I2C")

    add_heading(doc, "Backend and Infrastructure Tools", level=2)
    add_bullet(doc, "Eclipse Mosquitto 2.0 — MQTT broker running in Docker. "
               "Configured for anonymous access, persistence enabled.", "Mosquitto")
    add_bullet(doc, "InfluxDB 2.7 — time-series database. Events arrive via Telegraf "
               "and are stored in the 'attacks' bucket with 365-day retention.", "InfluxDB")
    add_bullet(doc, "Telegraf 1.30 — subscribes to honeypot/events/# on Mosquitto, parses "
               "JSON, and writes line-protocol points to InfluxDB.", "Telegraf")
    add_bullet(doc, "Grafana 10.4.0 — dashboard server. Two pre-provisioned dashboards "
               "(honeypot_overview.json, honeypot_patterns.json) auto-refresh every 10 s.", "Grafana")
    add_bullet(doc, "FastAPI — Python microservice (backend/classifier/) that polls InfluxDB "
               "every 60 s, runs heuristic + HMM + anomaly classification, and writes "
               "enriched pattern labels back.", "FastAPI classifier")
    add_bullet(doc, "Docker Compose — orchestrates all five backend containers on a shared "
               "honeypot_net bridge network with health checks.", "Docker Compose")

    add_heading(doc, "Development and Testing Tools", level=2)
    add_bullet(doc, "read_serial.py — Python script that reads the ESP32 serial port and "
               "saves output to serial_output.log for offline analysis.", "Serial monitor script")
    add_bullet(doc, "scripts/simulate_attacks.py — sends synthetic MQTT events covering all "
               "47 attack patterns, allowing the full backend pipeline to be exercised "
               "without physical hardware.", "Attack simulator")
    add_bullet(doc, "scripts/check_pipeline.sh — shell health-check script that verifies "
               "Mosquitto, InfluxDB, Telegraf, and Grafana are all healthy and data is "
               "flowing end-to-end.", "Pipeline health check")
    add_bullet(doc, "scripts/benchmark.py — measures event-processing throughput of the "
               "classifier microservice.", "Benchmark script")
    add_bullet(doc, "VS Code with the PlatformIO IDE extension — provides IntelliSense, "
               "build tasks, and serial monitor integration. "
               "The .vscode/c_cpp_properties.json configures the include path for ESP-IDF headers.",
               "VS Code / PlatformIO IDE")
    add_bullet(doc, "Git — version control. Secrets (include/secrets.h, backend/.env) are "
               "listed in .gitignore and never committed.", "Git")

    doc.add_paragraph()

    # ── Q4 ───────────────────────────────────────────────────────────────────
    add_heading(doc, "4. Real-Time System (RTS) — Course Project")

    add_para(doc,
        "The ESP32 honeypot is a real-time embedded system: it must respond to incoming "
        "network connections within bounded time, regardless of concurrent activity on "
        "other ports. Missing a connection window (TCP backlog overflow) or responding "
        "too slowly (attacker disconnects before the banner is sent) reduces the fidelity "
        "of captured data.")

    add_heading(doc, "Timing Constraints", level=2)
    add_bullet(doc,
        "TCP connections on all four ports are accepted within the same main loop iteration. "
        "The AsyncTCP library delivers connection callbacks within the lwIP interrupt context, "
        "ensuring <1 ms latency from SYN arrival to handler invocation.",
        "Connection acceptance")
    add_bullet(doc,
        "auth_success, command, and exploit events are published to MQTT immediately "
        "(not deferred to the 30-second batch). The deadline is soft — a publish failure "
        "triggers exponential backoff rather than a hard fault — but the design intent "
        "is that high-value events arrive at InfluxDB within 2–3 seconds.",
        "Immediate-publish deadline")
    add_bullet(doc,
        "The heartbeat fires every 30 seconds (EventLogger::loop compares millis() against "
        "_lastFlush). This is a periodic task with a 30 s period and a relaxed deadline: "
        "a missed heartbeat delays data by one period but does not corrupt system state.",
        "30 s heartbeat period")
    add_bullet(doc,
        "The Wi-Fi reconnect watchdog in WiFiManager::reconnectLoop() attempts reconnection "
        "if the link drops. The 15-second reconnect timeout was chosen to balance recovery "
        "speed against the risk of spamming association requests.",
        "Wi-Fi watchdog")

    add_heading(doc, "Network Reachability and NAT Constraint", level=2)
    add_para(doc,
        "The ESP32 obtains a private IP address (e.g. 192.168.x.x) from the DHCP server "
        "of the connected Wi-Fi access point or hotspot. This address is in the RFC 1918 "
        "private address space and is not routable from the public internet.")
    add_para(doc,
        "The fundamental reason is Network Address Translation (NAT): the router/hotspot "
        "presents a single public IP to the internet and maps outbound connections from "
        "private IPs to that public IP using port-state tracking. There is no pre-existing "
        "NAT entry for unsolicited inbound packets destined for the ESP32's private IP, "
        "so the router drops them before they ever reach the device.")
    add_para(doc,
        "For the honeypot to be reachable by external attackers, one of the following "
        "must be true:")
    add_bullet(doc,
        "The attacker is on the same local network (LAN) as the ESP32, so no NAT traversal "
        "is required.",
        "Same LAN")
    add_bullet(doc,
        "Port forwarding rules are configured on the router to map a public port to the "
        "ESP32's private IP on ports 22, 23, 80, and 1883. Without this, the router "
        "discards inbound TCP SYN packets because there is no NAT mapping for an "
        "unsolicited connection to a private address.",
        "Port forwarding on router")
    add_para(doc,
        "In the current lab deployment the ESP32 connects to a mobile hotspot (ZONGLE). "
        "The hotspot performs NAT between its LTE uplink and the Wi-Fi segment. "
        "External attack traffic is therefore not reachable unless port forwarding is "
        "configured on the carrier-grade NAT (CGNAT) of the mobile network, which is "
        "generally not possible without ISP cooperation. The honeypot is therefore "
        "effective only for capturing attacks from the same local network or from "
        "a host that is port-forwarded through a fixed broadband router.")

    add_heading(doc, "FreeRTOS Underpinning", level=2)
    add_para(doc,
        "The Arduino framework on ESP32 runs on top of FreeRTOS. The setup() and loop() "
        "functions execute in the app_main task (priority 1). AsyncTCP callbacks execute "
        "in a dedicated async_tcp task (priority configurable, default higher than app_main) "
        "created by the library, which is why connection handling is not blocked by "
        "a slow loop() iteration. This gives the system real-time responsiveness for "
        "I/O-bound events without explicit task creation in application code.")

    doc.add_paragraph()

    # ── Q5 ───────────────────────────────────────────────────────────────────
    add_heading(doc, "5. Real-Time Scheduling — Implementation and Course Project")

    add_heading(doc, "Scheduling Model Used: Cyclic Executive (Table-Driven)", level=2)
    add_para(doc,
        "The main loop() function is a hand-written cyclic executive. It calls each "
        "service handler in a fixed, repeating sequence:")

    code_text = (
        "void loop() {\n"
        "    WiFiManager::reconnectLoop();   // slot 1\n"
        "    MQTTService::loop();            // slot 2\n"
        "    TelnetHoneypot::loop();         // slot 3\n"
        "    SSHHoneypot::loop();            // slot 4\n"
        "    TLSHoneypot::loop();            // slot 5\n"
        "    EventLogger::loop();            // slot 6 — also fires 30s heartbeat\n"
        "    // LED update, periodic serial print ...\n"
        "}"
    )
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)

    add_para(doc,
        "This matches the cyclic executive pattern: tasks are arranged in a table "
        "(the source-code order), each executes non-preemptively in its assigned slot, "
        "and the major cycle repeats every ~1–5 ms depending on how many connections "
        "are active. There is no OS scheduler involved at the application level; "
        "timing discipline is enforced by each slot running to completion quickly.")

    add_heading(doc, "Why Not RMA or EDF?", level=2)
    add_para(doc,
        "Rate-Monotonic Analysis (RMA) and Earliest-Deadline-First (EDF) are applicable "
        "when tasks have fixed, known periods and deadlines that must be guaranteed. "
        "On this honeypot:")
    add_bullet(doc,
        "Network events are aperiodic and bursty — an attacker may send 100 lines in one "
        "second or nothing for an hour. A static priority assignment (RMA) would be "
        "under-utilised in the quiet case and offer no benefit in the burst case.",
        "Aperiodic arrivals")
    add_bullet(doc,
        "The only periodic task is the 30-second EventLogger heartbeat. A full EDF "
        "scheduler for one periodic task would add unnecessary overhead.",
        "Single periodic task")
    add_bullet(doc,
        "The AsyncTCP library already handles deadline-sensitive I/O (connection timeout "
        "CONFIG_ASYNC_TCP_MAX_ACK_TIME = 5000 ms) in its own higher-priority FreeRTOS task, "
        "so the application layer does not need a hard scheduler.",
        "I/O handled at higher priority")

    add_heading(doc, "Heartbeat as a Periodic Task — Informal RMA Analysis", level=2)
    add_para(doc,
        "If the heartbeat were treated as the sole periodic task with period T = 30 s, "
        "its utilisation is U = C/T where C is the execution time of flushQueue(). "
        "For 50 queued events, each requiring a ~2 ms MQTT publish, C ≈ 100 ms. "
        "U = 0.1 / 30 ≈ 0.003 (0.3%). This is well within the RMA schedulability "
        "bound of U ≤ n(2^(1/n) − 1) = 1.0 for n = 1 task. The heartbeat deadline "
        "is always met by a comfortable margin.")

    add_heading(doc, "MLF / LLF / LST — Not Applicable", level=2)
    add_para(doc,
        "Minimum Laxity First, Least Laxity First, and Least Slack Time are dynamic "
        "priority algorithms intended for systems with multiple concurrent tasks competing "
        "for a shared CPU. Because all honeypot service handlers complete within "
        "microseconds to low milliseconds and the only periodic task (heartbeat) has "
        "negligible utilisation, there is no scheduling conflict for MLF/LLF/LST to "
        "resolve. Applying them here would add context-switch overhead with no benefit.")

    add_heading(doc, "DMA — Not Used", level=2)
    add_para(doc,
        "Direct Memory Access is relevant when large data transfers must occur without "
        "CPU involvement (e.g. ADC streaming, SPI DMA). The honeypot's data transfers "
        "are small (< 512 bytes per event) and latency-tolerant; DMA is not used and "
        "is not required by the workload.")

    # Save
    out_path = os.path.join(OUT_DIR, "CAT2_Honeypot_Report.docx")
    doc.save(out_path)
    print(f"\nDocument saved → {out_path}")
    return out_path


if __name__ == "__main__":
    print("Generating diagrams...")
    fsm  = draw_fsm()
    cdfg = draw_cdfg()
    des  = draw_des()
    print("Building document...")
    out  = build_doc(fsm, cdfg, des)
    print("Done.")
